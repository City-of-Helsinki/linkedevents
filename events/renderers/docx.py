import collections
import datetime
import io

from django.utils.html import strip_tags
from django.utils.text import slugify
from django.utils.timezone import localtime
from docx import Document
from rest_framework import renderers
from rest_framework.utils.serializer_helpers import ReturnDict
from events.utils import parse_time


def get_any_language(dictionary, default='fi', language_codes=None):
    if not language_codes:
        # Default order for when a language isn't found.
        language_codes = [default, 'fi', 'sv', 'en']

    for language_code in language_codes:
        content = dictionary.get(language_code)
        if content:
            return content
    return ''


def get_description(event):
    if event['short_description']:
        return event['short_description']
    elif event['description']:
        return event['description']
    else:
        return {}


def get_price(raw_event):
    '''Returns either None or a dict with the price in each language.'''
    try:
        return raw_event['offers'][0]['price']
    except (IndexError, KeyError):
        return None


class EventParser:
    def __init__(self):
        self.earliest_date = None
        self.latest_date = None

    def parse_event(self, raw_event):
        '''
        Turns a raw event into a dictionary with only needed fields.
        Keeps track of earliest and latest times encountered.
        '''
        start_time = localtime(raw_event['start_time_obj'])
        end_time = localtime(raw_event['end_time_obj'])

        if self.earliest_date is None or start_time.date() < self.earliest_date:
            self.earliest_date = start_time.date()

        if self.latest_date is None or end_time.date() > self.latest_date:
            self.latest_date = start_time.date()

        price = get_price(raw_event)
        if price is not None:
            price = get_any_language(price)
        else:
            price = ''

        return {
            'name': get_any_language(raw_event['name']),
            'description': strip_tags(
                get_any_language(get_description(raw_event))
            ),
            'location': raw_event['location'],
            'start_time': start_time,
            'end_time': end_time,
            'price': price,
        }


def group_by_location(events):
    # We want to put events in the same location under the same headline,
    # so we create a dictionary where the places are the keys and the events
    # are the values.
    events_by_location = collections.defaultdict(list)
    for event in events:
        events_by_location[event['location']].append(event)
    return events_by_location


class DateRange:
    # Getting these by changing localization doesn't seem to work,
    # ideally you would use start.strftime(short_date + ' %A') instead.
    weekdays = [
        'maanantai',
        'tiistai',
        'keskiviikko',
        'torstai',
        'perjantai',
        'lauantai',
        'sunnuntai',
    ]

    def __init__(self, start, end, previous=None):
        self.start = start
        self.end = end
        self.previous = previous

    def __str__(self):
        start = self.start
        end = self.end

        long_date = '%-d.%-m.%Y'
        short_date = long_date

        if self.previous is not None:
            if self.previous.start.year == self.start.year:
                short_date = '%-d.%-m.'

        # A single day with locale weekday name, e.g. '1.1. maanantai'
        if start == end:
            return start.strftime(short_date + ' ' + self.weekdays[start.weekday()])

        # Multiple days, e.g. 1.1.-2.1.
        if start.year == end.year:
            return '%s-%s' % (
                start.strftime(short_date),
                end.strftime(short_date),
            )

        # Multiple years, e.g. 1.1.2017-31.12.2019
        return '%s-%s' % (
            start.strftime(long_date),
            end.strftime(long_date),
        )

    def __lt__(self, other):
        return self.start < other.start

    def __eq__(self, other):
        return self.start == other.start and self.end == other.end

    def __hash__(self):
        return hash(str(self.start) + str(self.end))


def group_by_date(events):
    # We want events on the same date to be under the same day headline,
    # and events that span multiple dates to be under their own headlines.
    events = sorted(events, key=lambda e: e['start_time'])

    dates = collections.defaultdict(list)

    previous_daterange = None
    for event in events:
        date_range = DateRange(
            event['start_time'].date(),
            event['end_time'].date(),
            previous=previous_daterange,
        )
        previous_daterange = date_range

        dates[date_range].append(event)

    dates = collections.OrderedDict(sorted(dates.items()))

    return dates


class DOCXRenderer(renderers.BaseRenderer):
    media_type = (
        'application/vnd.openxmlformats-officedocument'
        '.wordprocessingml.document'
    )
    format = 'docx'
    charset = None
    render_style = 'binary'

    def get_document(self):
        # This is here so that the document class can be swapped out to
        # facilitate testing. Due to some threading issues this returns an
        # object and not a class
        return Document()

    def render(self, data, media_type=None, renderer_context=None):
        document = self.get_document()
        parsed_events = []

        query_params = renderer_context['request'].query_params

        event_parser = EventParser()
        if type(data) is ReturnDict:
            # Support the single event endpoint just because we can
            data = [data]

        first_location = data[0]['location']
        for raw_event in data:
            parsed_events.append(event_parser.parse_event(raw_event))

        # This is here to allow for this to be expanded to include multiple
        # locations in the future.
        locations = group_by_location(parsed_events)

        for location in locations:
            locations[location] = group_by_date(locations[location])

        # We need to get the daterange for the entire document, which is
        # determined by either the query or the actual events.
        start = query_params.get('start')
        end = query_params.get('end')

        if start is None:
            query_start_date = event_parser.earliest_date
        else:
            query_start_date = parse_time(start, True)[0]

        if end is None:
            query_end_date = event_parser.latest_date
        else:
            query_end_date = parse_time(end, False)[0]

        total_date_range = DateRange(query_start_date, query_end_date)

        midnight = datetime.time(0, 0)
        for location, dateranges in locations.items():
            document.add_heading(str(location), 0)
            document.add_paragraph(str(total_date_range))

            for daterange, events in dateranges.items():
                document.add_heading(str(daterange), 1)

                for event in events:

                    # This is here to prevent 00:00-00:00 from being shown.
                    start_time = event['start_time']
                    end_time = event['end_time']

                    if start_time.time() == end_time.time() == midnight:
                        document.add_heading(event['name'], 2)
                    else:
                        document.add_heading(
                            '%s-%s %s' % (
                                event['start_time'].strftime('%H:%M'),
                                event['end_time'].strftime('%H:%M'),
                                event['name'],
                            ),
                            2,
                        )

                    document.add_paragraph(event['description'])
                    if event['price']:
                        document.add_paragraph(event['price'])

        filename = '%s-%s-%s.docx' % (
            slugify(first_location.name),
            query_start_date.strftime('%Y%m%d'),
            query_end_date.strftime('%Y%m%d'),
        )

        renderer_context['response']['Content-Disposition'] = (
            'attachment; filename=%s' % filename
        )

        output = io.BytesIO()
        document.save(output)

        return output.getvalue()
